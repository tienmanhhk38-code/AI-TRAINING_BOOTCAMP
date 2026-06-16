from csv import reader as csv_reader
from io import BytesIO, StringIO
from pathlib import Path

SUPPORTED_EXTENSIONS = {".txt", ".md", ".csv", ".pdf", ".docx", ".xlsx", ".xlsm"}

class DocumentLoadError(ValueError):
    pass

def extension_for(filename):
    return Path(filename).suffix.lower()

def decode_text(content):
    for encoding in ("utf-8-sig", "utf-8", "cp1258", "latin-1"):
        try:
            return content.decode(encoding)
        except UnicodeDecodeError:
            continue
    raise DocumentLoadError("Cannot decode text file")

def load_document_text(filename, content):
    extension = extension_for(filename)
    if extension not in SUPPORTED_EXTENSIONS:
        supported = ", ".join(sorted(SUPPORTED_EXTENSIONS))
        raise DocumentLoadError(f"Unsupported file type {extension}. Supported: {supported}")

    if extension in {".txt", ".md"}:
        return decode_text(content).strip()

    if extension == ".csv":
        return load_csv_text(content)

    if extension == ".pdf":
        return load_pdf_text(content)

    if extension == ".docx":
        return load_docx_text(content)

    if extension in {".xlsx", ".xlsm"}:
        return load_excel_text(content)

    raise DocumentLoadError(f"Unsupported file type {extension}")

def load_csv_text(content):
    text = decode_text(content)
    rows = []
    for row in csv_reader(StringIO(text)):
        cleaned = [cell.strip() for cell in row if cell.strip()]
        if cleaned:
            rows.append(" | ".join(cleaned))
    return "\n".join(rows).strip()

def load_pdf_text(content):
    try:
        from pypdf import PdfReader
    except ImportError as error:
        raise DocumentLoadError("Missing pypdf dependency") from error

    try:
        pdf = PdfReader(BytesIO(content))
        pages = [page.extract_text() or "" for page in pdf.pages]
    except Exception as error:
        raise DocumentLoadError(f"Cannot read PDF: {type(error).__name__}") from error

    return "\n".join(page.strip() for page in pages if page.strip()).strip()

def load_docx_text(content):
    try:
        from docx import Document
    except ImportError as error:
        raise DocumentLoadError("Missing python-docx dependency") from error

    try:
        document = Document(BytesIO(content))
    except Exception as error:
        raise DocumentLoadError(f"Cannot read DOCX: {type(error).__name__}") from error

    paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    table_rows = []
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                table_rows.append(" | ".join(cells))

    return "\n".join(paragraphs + table_rows).strip()

def load_excel_text(content):
    try:
        from openpyxl import load_workbook
    except ImportError as error:
        raise DocumentLoadError("Missing openpyxl dependency") from error

    try:
        workbook = load_workbook(BytesIO(content), read_only=True, data_only=True)
    except Exception as error:
        raise DocumentLoadError(f"Cannot read Excel file: {type(error).__name__}") from error

    lines = []
    for worksheet in workbook.worksheets:
        lines.append(f"Sheet: {worksheet.title}")
        for row in worksheet.iter_rows(values_only=True):
            values = [str(value).strip() for value in row if value is not None and str(value).strip()]
            if values:
                lines.append(" | ".join(values))

    return "\n".join(lines).strip()
